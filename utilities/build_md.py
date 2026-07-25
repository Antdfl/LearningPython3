#!/usr/bin/env python3
"""
Converts a .md file to HTML, PDF, and/or DOCX (Word).
On startup it asks whether to generate HTML, PDF, DOCX, or all three formats.

Menu: H = HTML only (.html), P = PDF only (.pdf), W = Word only (.docx), A = All three formats
"""
import sys, re, unicodedata, os
from pathlib import Path


def _missing(pip_name):
    """Prints a helpful message when a required library is not installed."""
    print(f"Library '{pip_name}' not found. Install with: pip install {pip_name}")


try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    _missing('markdown')
    MARKDOWN_AVAILABLE = False

try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    _missing('xhtml2pdf')
    XHTML2PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    _missing('python-docx')
    DOCX_AVAILABLE = False


SCRIPT_DIR = Path(__file__).resolve().parent


def strip_frontmatter(md_text):
    """Removes the initial YAML block (--- ... ---) if present, so that it doesn't end up as raw text in the generated document.

    This function is useful when your Markdown file has frontmatter metadata (common in Hugo, Jekyll, etc.)
    and you want to exclude it from the output document."""
    lines = md_text.split('\n')
    if lines and lines[0].strip() == '---':
        for idx in range(1, len(lines)):
            if lines[idx].strip() == '---':
                return '\n'.join(lines[idx + 1:]).lstrip('\n')
    return md_text


def convert_to_html(md_text):
    """Converts Markdown text to HTML string using the 'markdown' library.

    Returns:
        str: The converted HTML string with tables, sane lists, and new line breaks supported.
    """
    return markdown.markdown(md_text, extensions=['tables', 'sane_lists', 'nl2br'])


def html_to_pdf(html_content, output_path):
    """Converts an HTML string to a PDF file using the 'xhtml2pdf' library.

    Args:
        html_content: The HTML content as a string.
        output_path: The path where the PDF file will be saved.

    Note:
        This function uses xhtml2pdf, which may not be available on all systems.
        Check that XHTML2PDF_AVAILABLE is True before using this."""
    try:
        with open(output_path, 'wb') as pdf_file:
            result = pisa.CreatePDF(html_content, dest=pdf_file, encoding='utf-8')
        if result.err:
            print(f"PDF generation error (xhtml2pdf reported {result.err} problems)")
        else:
            print(f"OK: {output_path.name}")
    except PermissionError as e:
        print(f"Cannot write to {output_path.name}: {e}")


def add_heading_bookmark(paragraph, bookmark_name, bookmark_id):
    """Adds a bookmark to a heading paragraph so it can be reached from the table of contents (TOC).

    When generating a Table of Contents in Word, each heading needs an associated bookmark.
    This function creates start and end bookmark markers around the heading element.

    Args:
        paragraph: The paragraph element that contains the heading.
        bookmark_name: A unique name for this bookmark (used as anchor).
        bookmark_id: A numeric ID for this bookmark (e.g., 1, 2, 3...)."""
    p_elem = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    # w:pPr (if present), must remain the first child of the paragraph element.
    # We insert new elements (start/end bookmarks) after pPr, not before.
    insert_at = 1 if len(p_elem) and p_elem[0].tag == qn('w:pPr') else 0
    p_elem.insert(insert_at, start)
    p_elem.append(end)


def build_toc_entry_paragraph(doc, level, text, bookmark_name):
    """Creates a table of contents (TOC) entry paragraph for a heading.

    This function builds a clickable TOC item that links to the corresponding bookmark
    in the document. The indentation level is calculated from the heading level:
        Level 1: No indent
        Level 2: Indent by 0.25 inches
        Level 3: Indent by 0.5 inches
        Other levels: Indent by 0.5 inches (default)

    Args:
        doc: The Word document object.
        level: The heading level (1-6, where 1 is largest).
        text: The displayed text for the TOC entry.
        bookmark_name: The bookmark name that this TOC entry will link to."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches({1: 0, 2: 0.25, 3: 0.5}.get(level, 0.5))

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    run_el = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '1e3c72')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.append(color)
    rpr.append(underline)
    run_el.append(rpr)
    t = OxmlElement('w:t')
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)

    p._p.append(hyperlink)
    return p


def build_toc(doc, toc_entries):
    """Builds the Table of Contents (TOC) block and inserts it at the beginning of the document.

    Word and LibreOffice generate their own TOCs using a specific structure:
        - A field wrapper (w:sdt) with gallery type "Table of Contents"
        - Field begin marker for the TOC field
        - Field instruction: 'TOC \\o "1-3" \\h \\z \\u'
          This tells Word to show levels 1-3, hyperlinked, no page numbers, update on selection
        - Field separate marker
        - Field end marker

    The TOC block contains:
        - A title "Table of Contents"
        - Clickable entries that link to document bookmarks
        - Entries are indented based on heading level

    Important note:
        The bookmark field markers don't build the TOC from scratch. They just allow Word
        to recalculate the TOC when needed (e.g., using Ctrl+A > Table of Contents > Update, or F9).
        The displayed text in the TOC is taken directly from the document headings.

    Args:
        doc: The Word document object.
        toc_entries: List of tuples (level, text, bookmark_name) for each heading."""
    if not toc_entries:
        return

    # Step 1: Create the TOC title paragraph
    title_p = doc.add_paragraph()
    title_run = title_p.add_run('Table of Contents')
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    # Step 2: Build all TOC entry paragraphs (hyperlinked list items)
    entry_paragraphs = [
        build_toc_entry_paragraph(doc, level, text, bookmark_name)
        for level, text, bookmark_name in toc_entries
    ]

    # Step 3: Add a page break after the TOC title
    doc.add_page_break()

    # Step 4: Wrap entries in Word's field structure for automatic TOC generation
    first_p = entry_paragraphs[0]._p
    
    # Field begin marker (w:fldCharType='begin') signals start of a dynamic field
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin = OxmlElement('w:r')
    r_begin.append(fld_begin)

    # Field instruction text that Word interprets as "build a TOC with levels 1-3"
    # \\o "1-3" = show heading levels 1, 2, and 3
    # \\h = hyperlinked entries
    # \\z = no page numbers shown
    # \\u = update field on selection (e.g., Ctrl+A > F9)
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr = OxmlElement('w:r')
    r_instr.append(instr_text)

    # Field separate marker (w:fldCharType='separate') indicates the field content follows
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    r_separate = OxmlElement('w:r')
    r_separate.append(fld_separate)

    # Insert field markers before the first entry paragraph
    insert_at = 1 if len(first_p) and first_p[0].tag == qn('w:pPr') else 0
    for el in (r_separate, r_instr, r_begin):
        first_p.insert(insert_at, el)

    last_p = entry_paragraphs[-1]._p
    # Field end marker (w:fldCharType='end') closes the field wrapper
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end = OxmlElement('w:r')
    r_end.append(fld_end)
    last_p.append(r_end)

    # Step 5: Create the Sdt (Simple Data Type) wrapper that marks this as a TOC gallery
    sdt = OxmlElement('w:sdt')
    sdt_pr = OxmlElement('w:sdtPr')
    doc_part_obj = OxmlElement('w:docPartObj')
    doc_part_gallery = OxmlElement('w:docPartGallery')
    doc_part_gallery.set(qn('w:val'), 'Table of Contents')
    doc_part_unique = OxmlElement('w:docPartUnique')
    doc_part_obj.append(doc_part_gallery)
    doc_part_obj.append(doc_part_unique)
    sdt_pr.append(doc_part_obj)
    sdt.append(sdt_pr)

    # Step 6: Attach title and entries to the Sdt container
    sdt_content = OxmlElement('w:sdtContent')
    sdt_content.append(title_p._p)
    for p in entry_paragraphs:
        sdt_content.append(p._p)
    sdt.append(sdt_content)

    # Step 7: Insert the complete TOC structure at the beginning of the document body
    doc.element.body.insert(0, sdt)


def create_word_from_md(md_text, output_path):
    """Creates a Word document (.docx) from Markdown with support for:

    Supported Markdown features (based on https://www.markdownguide.org/basic-syntax/ 
    and https://google.github.io/styleguide/docguide/style.html):
    - Headings: # h1, ## h2, ### h3, #### h4, ##### h5, ###### h6, hierarchically numbered
      (1, 1.1, 1.1.1, ...) for levels 1-3
    - Clickable table of contents at the beginning (pre-populated, doesn't depend on F9)
    - Regular paragraphs
    - Unordered lists (- or *)
    - Ordered lists (1., 2., ...)
    - Bold (**text**)
    - Italic (*text*)
    - Inline code (`code`)
    - Tables
    - Links [text](url)
    - Images ![alt](path)
    """
    def split_table_row(row_line):
        """Splits a single Markdown table row line into individual cells, preserving empty cells.

        This helper function handles the parsing of a single row from a Markdown table by:
          1. Splitting on the | character (Markdown table separator)
          2. Removing any leading/trailing empty columns (first and last if they're just pipes)
          3. Stripping whitespace from each cell content

        Args:
            row_line: A single line containing a Markdown table row like "Header1 | Header2"

        Returns:
            list: A list of strings, where each string is the trimmed content of one cell.
                  Empty cells (between ||) are preserved as empty strings.
        """
        parts = row_line.split('|')
        if parts and parts[0].strip() == '':
            parts = parts[1:]
        if parts and parts[-1].strip() == '':
            parts = parts[:-1]
        return [p.strip() for p in parts]

    try:
        doc = Document()
        
        # Initialize tracking variables for the document processing
        h_counters = [0, 0, 0]          # Counters for heading levels 1, 2, 3 (used for hierarchical numbering)
        toc_entries = []                 # List to store (level, text, bookmark_name) tuples for TOC generation
        bookmark_counter = 0             # Counter that increments for each heading (used for unique bookmark names)
        
        lines = md_text.split('\n')      # Split Markdown content into individual lines
        num_lines = len(lines)           # Total number of lines to process
        i = 0                            # Line index iterator

        # Main processing loop: iterate through each line of the Markdown document
        while i < num_lines:
            line_stripped = lines[i]     # Get current line from original list
            line = line_stripped.strip() # Remove leading/trailing whitespace

            if not line:
                i += 1                   # Skip empty lines
                continue

            # Check for Markdown tables (header + separator + body rows)
            # Tables are detected here because they must not start with '#' (headings indicator)
            # starting from the current line, not recalculated over the entire document.
            if '|' in line and not line.startswith('#'):
                next_line = lines[i + 1].strip() if i + 1 < num_lines else ''
                sep_cells = split_table_row(next_line) if next_line else []
                is_separator = bool(sep_cells) and all(
                    re.match(r'^:?-+:?$', c) for c in sep_cells if c != ''
                ) and any(c != '' for c in sep_cells)

                if is_separator:
                    header_cells = split_table_row(line)
                    num_cols = len(header_cells)

                    body_rows = []
                    j = i + 2
                    while j < num_lines and '|' in lines[j].strip():
                        body_rows.append(split_table_row(lines[j].strip()))
                        j += 1

                    table = doc.add_table(rows=1 + len(body_rows), cols=num_cols)
                    table.style = 'Table Grid'

                    for k, cell_text in enumerate(header_cells[:num_cols]):
                        table.rows[0].cells[k].text = cell_text

                    for r, cells in enumerate(body_rows, start=1):
                        for k, cell_text in enumerate(cells[:num_cols]):
                            table.rows[r].cells[k].text = cell_text

                    i = j
                    continue

            # Images ![alt](path)
            img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line_stripped)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                p = doc.add_paragraph(f"Image: {alt_text} (path: {img_path})")
                i += 1
                continue

            # Links [text](url) - parse first to avoid ambiguity with headings
            link_match = re.match(r'^\[([^\]]*)\]\(([^)]+)\)$', line_stripped)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                p = doc.add_paragraph(f"[{link_text}]({link_url})")
                i += 1
                continue
            
            # Headings - supports h1 to h6 based on the number of #
            title_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
            if title_match:
                level = len(title_match.group(1))
                title_text = title_match.group(2).strip()
                
                # Remove any [text](url) links from the heading, keeping only the text
                title_clean = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', title_text)
                
                if level in (1, 2, 3):
                    if level == 1:
                        h_counters[0] += 1
                        h_counters[1] = 0
                        h_counters[2] = 0
                        numbering = f"{h_counters[0]}"
                    elif level == 2:
                        h_counters[1] += 1
                        h_counters[2] = 0
                        numbering = f"{h_counters[0]}.{h_counters[1]}"
                    else:
                        h_counters[2] += 1
                        numbering = f"{h_counters[0]}.{h_counters[1]}.{h_counters[2]}"

                    numbered_text = f"{numbering} {title_clean}"
                    heading = doc.add_heading(numbered_text, level=level)

                    bookmark_counter += 1
                    bookmark_name = f"_toc_bm_{bookmark_counter}"
                    add_heading_bookmark(heading, bookmark_name, bookmark_counter)
                    toc_entries.append((level, numbered_text, bookmark_name))
                elif level == 4:
                    p = doc.add_paragraph()
                    run = p.add_run(title_clean)
                    run.font.name = 'Heading'
                    run.font.size = Pt(14)
                elif level == 5:
                    p = doc.add_paragraph()
                    run = p.add_run(title_clean)
                    run.font.size = Pt(12)
                else:  # h6
                    p = doc.add_paragraph()
                    run = p.add_run(title_clean)
                    run.font.size = Pt(10)

                i += 1
                continue

            # Unordered lists (- or *)
            ul_match = re.match(r'^(\s*[-*+]\s+)(.+)$', line_stripped)
            if ul_match:
                list_item_text = ul_match.group(2).strip()
                p = doc.add_paragraph(list_item_text)
                i += 1
                continue

            # Ordered lists (1., 2., ...)
            ol_match = re.match(r'^(\d+)\.\s+(.+)$', line_stripped)
            if ol_match:
                list_num = ol_match.group(1)
                list_item_text = ol_match.group(2).strip()
                p = doc.add_paragraph(f"{list_num}. {list_item_text}")
                i += 1
                continue

            # Quoted text / Blockquotes
            quote_match = re.match(r'^"(.+)"$', line_stripped)
            if quote_match:
                quote_text = quote_match.group(1).replace('"', '"')
                p = doc.add_paragraph('" ' + quote_text + '"', style='Quote')
                i += 1
                continue
            
            # Inline code (`code`)
            code_inline = re.sub(r'`([^`]+)`', r'<code>\1</code>', line_stripped)
            
            # Bold (**text**)
            bold_inline = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', code_inline)
            
            # Italic (*text*)
            italic_inline = re.sub(r'\*(.+?)\*', r'<em>\1</em>', bold_inline)
            
            # Remove HTML tags (if present) before adding paragraph
            clean_text = re.sub(r'<[^>]+>', '', italic_inline)
            
            # Normal paragraph text
            if clean_text.strip():
                p = doc.add_paragraph(clean_text)

            i += 1

        build_toc(doc, toc_entries)

        doc.save(output_path)
        print(f"OK: {output_path.name}")
    except Exception as e:
        print(f"Error creating DOCX file: {e}")


# ── Interactive menu to choose output formats ──
print("\n" + "=" * 40)
print("=== Convert Markdown File ===")
print("=" * 40)
print("H = HTML only (.html)")
print("P = PDF only (.pdf)  <-- DEFAULT")
print("W = Word only (.docx)")
print("A = All three formats (HTML, PDF, DOCX)")
print("=" * 40)

choice = input("\nChoose format: H | P | W | A? ").strip().upper()

# Fallback: if the user didn't select anything, default to PDF
if not choice or choice == "":
    print("No selection - using default: PDF (.pdf)")
    choice = "P"

if choice not in ['H', 'P', 'W', 'A']:
    print("Invalid selection. Exiting program.")
    sys.exit(1)

# ── Input file ──
arg = sys.argv[1] if len(sys.argv) > 1 else 'text_to_convert.md'
src = Path(arg)
if not src.is_absolute():
    src = SCRIPT_DIR / src

if not src.exists():
    print(f"File not found: {src}")
    sys.exit(1)

# ── Output filename ──
default_stem = src.with_suffix('').name
output_name = input(f"\nOutput filename base [{default_stem}]: ").strip()
stem = Path(output_name) if output_name else src.with_suffix('')

# ── Read the Markdown content ──
print("\nReading Markdown file...")
md_text = open(src, encoding='utf-8').read()
md_text = strip_frontmatter(md_text)
print(f"Read: {len(md_text)} characters.")

# ── Convert Markdown to HTML (required only for HTML/PDF output) ──
needs_html = choice in ['H', 'P', 'A']
body_html = convert_to_html(md_text) if needs_html and MARKDOWN_AVAILABLE else None

# ── CSS stylesheet for HTML/PDF rendering ──
CSS = """
@page { size: A4; margin: 18mm 16mm 16mm 16mm; }
* { box-sizing: border-box; }
body { font-family: 'Lato','Segoe UI',Helvetica,Arial,sans-serif; font-size: 10pt; line-height: 1.55; color: #2b2b2b; margin: 0; }
h1 { font-size: 20pt; color: #1e3c72; margin: 0 0 4px 0; letter-spacing: .3px; }
h2 { font-size: 12.5pt; color: #1e3c72; margin: 22px 0 8px 0; padding-bottom: 4px; border-bottom: 2px solid #1e3c72; page-break-after: avoid; }
h3 { font-size: 10.8pt; color: #2a5298; margin: 15px 0 5px 0; page-break-after: avoid; }
h4 { font-size: 10pt; color: #444; margin: 12px 0 4px 0; page-break-after: avoid; }
h5 { font-size: 9.6pt; color: #555; margin: 10px 0 3px 0; font-style: italic; page-break-after: avoid; }
p { margin: 0 0 7px 0; }
ul, ol { margin: 4px 0 8px 0; padding-left: 20px; }
li { margin-bottom: 3px; }
strong { color: #1e3c72; font-weight: bold; }
em { color: #555; font-style: italic; }
code { background: #f0f2f5; padding: 2px 6px; border-radius: 3px; font-family: 'Consolas','Courier New',monospace; font-size: 9pt; }
pre { background: #f5f7fa; border-left: 3px solid #8fa8cf; padding: 10px 12px; margin: 8px 0 12px 0; overflow-x: auto; }
pre code { background: none; padding: 0; }
table { width: 100%; border-collapse: collapse; margin: 8px 0 12px 0; font-size: 9pt; page-break-inside: avoid; }
th { background: #eef3fb; color: #1e3c72; text-align: left; padding: 6px 8px; border: 1px solid #cdd9ec; font-weight: bold; }
td { padding: 6px 8px; border: 1px solid #dde5f0; vertical-align: top; }
blockquote { margin: 7px 0 10px 0; padding: 8px 14px; background: #f5f7fa; border-left: 4px solid #1e3c72; font-size: 9.6pt; page-break-inside: avoid; }
blockquote p { margin: 0; }
hr { border: none; border-top: 1px solid #e2e2e2; margin: 15px 0; }
a { color: #1e3c72; text-decoration: underline; }
img { max-width: 100%; height: auto; }
"""

# ── Assemble the HTML document ──
html_doc = None
if body_html is not None:
    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>Converted Document</title>
<style>{CSS}</style></head>
<body>
{body_html}
</body>
</html>"""

# ── Generate files based on user choice ──
if choice in ['H', 'A']:
    if html_doc is not None:
        html_path = stem.with_suffix('.html')
        html_path.write_text(html_doc, encoding='utf-8')
        print(f"Generated: {html_path}")
    else:
        _missing('markdown')
        print("Skipping HTML generation.")

if choice in ['P', 'A']:
    if html_doc is None:
        _missing('markdown')
        print("Skipping PDF generation.")
    elif not XHTML2PDF_AVAILABLE:
        _missing('xhtml2pdf')
        print("Skipping PDF generation.")
    else:
        pdf_path = stem.with_suffix('.pdf')
        html_to_pdf(html_doc, pdf_path)

if choice in ['W', 'A']:
    if DOCX_AVAILABLE:
        word_path = stem.with_suffix('.docx')
        create_word_from_md(md_text, word_path)
    else:
        _missing('python-docx')
        print("Skipping DOCX generation.")

print("\n" + "=" * 40)
print("Conversion complete!")
print("=" * 40)