#!/usr/bin/env python3
"""
utilities/build_md_toc.py

# Purpose: This module provides utilities for converting Markdown (.md) files into multiple document 
# formats (HTML, PDF, DOCX). It is designed to handle complex features like automatic Table of Contents (TOC) generation.

# Audience Note for Junior Programmers:
# The conversion process involves several technical workarounds and external library dependencies 
# (markdown, xhtml2pdf, python-docx). Understanding the *why* behind these functions 
# (e.g., why table widths must be explicitly set in CSS/XML) is crucial for maintenance.

# Functionality Overview:
# - Markdown to HTML: Converts markdown syntax to structured, styled HTML.
# - TOC Generation (HTML/PDF): Dynamically finds H1-H3 headings, creates unique anchors, and builds a navigable TOC block placed after the first major heading.
# - DOCX Generation: Uses advanced python-docx XML manipulation to create a Word document that contains proper internal field codes for an automatic TOC when opened in MS Word.

# Conversion Menu:
# H = HTML only (.html)
# P = PDF only (.pdf): Requires xhtml2pdf and relies on the generated HTML structure.
# W = Word only (.docx): Uses advanced DOCX XML manipulation for full functionality.
# A = All three formats: Executes all supported conversions sequentially.

# Dependencies (MUST be installed via pip): 
# - markdown: Core Markdown parser.
# - xhtml2pdf: For PDF conversion from HTML.
# - python-docx: For Word document creation and XML manipulation.
"""
import sys, re, unicodedata, os
from pathlib import Path


# ── Helper comments section ──


def _missing(pip_name: str) -> None:
    """
    [Helper Function] Outputs an actionable warning message if a required third-party library 
    is missing from the current Python environment. This function is critical for robust
    runtime checks, preventing unexpected crashes due to missing dependencies.

    Parameters:
        pip_name (str): The name of the missing library package (e.g., 'markdown').
    """
    # Use a simple f-string print statement for immediate feedback to the user/developer running the script.
    print(f"Libreria '{pip_name}' non trovata. Installala con: pip install {pip_name}")


try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    _missing('markdown')
    MARKDOWN_AVAILABLE = False

try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    _missing('xhtml2pdf')
    XHTML2PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    _missing('python-docx')
    DOCX_AVAILABLE = False


SCRIPT_DIR = Path(__file__).resolve().parent


def strip_frontmatter(md_text: str) -> str:
    """
    [Utility Function] Strips common YAML "Front Matter" blocks from the start of a Markdown file.

    The front matter is metadata (like author, date, or version) used by static site generators 
    (e.g., Jekyll, Hugo). Since this data is *outside* the primary narrative content, it must be removed 
    before parsing to prevent the parser from interpreting meta-data keys as regular text/headings.

    Parameters:
        md_text (str): The raw Markdown content string read from disk. This is expected to contain
            the '---' delimiters defining the front matter block.

    Returns:
        str: The cleaned Markdown text with the front matter block completely removed, maintaining 
             the original line endings and structure of the core document content.
    """
    lines = md_text.split('\n') # Splits the string into a list of lines for iterative checking.
    # Check if the first line is the start delimiter '---'.
    if lines and lines[0].strip() == '---':
        for idx in range(1, len(lines)):
            # Found the closing delimiter '---'. Everything after this point is content.
            if lines[idx].strip() == '---':
                return '\n'.join(lines[idx + 1:]).lstrip('\n')
    # If no front matter structure was found, return the original text unchanged.
    return md_text


def extract_frontmatter_fields(md_text):
    """
    [Utility Function] Reads simple flat 'key: value' pairs out of a leading
    YAML front matter block - just enough to recover metadata like 'title'
    and 'subtitle' for display, not a general-purpose YAML parser.

    Parameters:
        md_text (str): The raw Markdown content string, front matter still
            attached (call this before strip_frontmatter removes it).

    Returns:
        dict: Front matter keys mapped to their unquoted string values.
        Empty dict if the document has no front matter block.
    """
    fields = {}
    lines = md_text.split('\n')
    if not lines or lines[0].strip() != '---':
        return fields
    for line in lines[1:]:
        if line.strip() == '---':
            break
        match = re.match(r'^(\w[\w-]*)\s*:\s*(.+)$', line.strip())
        if match:
            key, value = match.group(1), match.group(2).strip()
            fields[key] = value.strip('"').strip("'")
    return fields


def strip_md_toc(md_text):
    """
    [Utility Function] Strips a hand-written Markdown Table of Contents (TOC) from the
    start of a document, so it does not duplicate the auto-generated TOC built later for
    HTML/PDF/DOCX output.

    Two TOC shapes are recognized, both searched for only near the top of the document
    (right after the title, so real content lists elsewhere are never touched):

    1. Heading-guarded: a heading like "## Contenuto" / "## Table of Contents" / "## Indice"
       followed by a list of anchor links (e.g. "[Link](#id)").
    2. Bare: a contiguous list of anchor links placed directly under the first H1 title,
       with no introducing heading at all (this is what tools like Pandoc/VS Code TOC
       extensions typically generate).

    A run of list items is only treated as a real TOC if the large majority of its items
    are anchor links, which avoids false positives on ordinary bullet lists.

    IMPORTANT: The first H1 heading (the document title) is never removed, only the TOC
    block that follows it.

    Args:
        md_text (str): The raw Markdown content string.

    Returns:
        str: The cleaned Markdown text with the TOC block removed, unchanged if none found.
    """
    lines = md_text.split('\n')

    if not lines:
        return md_text

    total = len(lines)
    # A hand-written TOC always sits near the very top of the document (right after the
    # title), so we only look there. This keeps ordinary content lists further down
    # (which may also happen to contain internal anchor links) untouched.
    search_limit = min(total, 400)

    # Heading text is anchored with \s*$ so a real title like "# Sommario del Progetto"
    # is never mistaken for a bare TOC heading like "# Sommario".
    heading_patterns = [
        r'^#\s*Contenuto\s*$',
        r'^#\s*Indice\s*$',
        r'^#\s*Table\s+of\s+Contents\s*$',
        r'^#\s*Sommario\s*$',
        r'^##\s*Contenuto\s*$',
        r'^##\s*Indice\s*$',
        r'^##\s*Table\s+of\s+Contents\s*$',
        r'^###\s*Contenuto\s*$',
    ]

    def is_list_item(s):
        return bool(re.match(r'^\s*[-*+]\s+', s))

    def is_anchor_list_item(s):
        return bool(re.match(r'^\s*[-*+]\s+.*\[.+?\]\(#.*?\)', s))

    # Locate the first H1 (document title); the TOC search starts right after it.
    title_idx = None
    for i in range(min(total, search_limit)):
        if re.match(r'^#\s+', lines[i].strip()):
            title_idx = i
            break

    scan_start = (title_idx + 1) if title_idx is not None else 0

    toc_start_idx = None
    toc_end_idx = None

    i = scan_start
    while i < search_limit:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        heading_line_idx = None
        if any(re.match(p, stripped) for p in heading_patterns):
            heading_line_idx = i
            j = i + 1
            while j < search_limit and not lines[j].strip():
                j += 1
            run_start = j
        elif is_anchor_list_item(stripped):
            run_start = i
        else:
            # First non-blank content after the title is neither a TOC heading nor a
            # TOC-like list item, so there is no hand-written TOC to strip.
            break

        # Collect the contiguous run of list items starting at run_start (a real blank
        # line ends the run - TOC entries are not blank-line separated).
        collected = []
        j = run_start
        while j < search_limit and is_list_item(lines[j].strip()):
            collected.append(j)
            j += 1

        if collected:
            anchor_count = sum(1 for k in collected if is_anchor_list_item(lines[k].strip()))
            if anchor_count / len(collected) >= 0.8:
                toc_start_idx = heading_line_idx if heading_line_idx is not None else collected[0]
                toc_end_idx = collected[-1]
        break

    if toc_start_idx is None or toc_end_idx is None:
        return md_text

    result_lines = lines[:toc_start_idx] + lines[toc_end_idx + 1:]
    return '\n'.join(result_lines)


_LEADING_HASH_RE = re.compile(r'^(#{1,6})(.*)$')
_FENCE_LINE_RE = re.compile(r'^(```|~~~)')


def escape_false_headings(md_text):
    """Neutralizes lines that start with 1-6 '#' characters but are not real
    Markdown ATX headings, before handing text to the 'markdown' library.

    The 'markdown' library's heading detection is more permissive than the
    CommonMark rule it is usually associated with: it treats ANY line
    starting with 1-6 '#' as a heading, even with no space afterwards (e.g.
    '#2 ...' used as a post number, or a lone '#PySpark' hashtag). A real
    ATX heading always has whitespace (or nothing) right after the hashes,
    so that is the signal used here to tell the two apart.

    Lines inside fenced code blocks (```/~~~) are left untouched, since a
    line like '#!/usr/bin/env python' or a '# comment' there is code, not a
    false heading, and must reach the renderer unescaped.

    Args:
        md_text (str): Raw Markdown content, before markdown.markdown().

    Returns:
        str: The same text with false-heading lines prefixed by a
            backslash, which 'markdown' renders as a literal '#'.
    """
    out_lines = []
    in_fence = False
    for line in md_text.split('\n'):
        if _FENCE_LINE_RE.match(line.lstrip()):
            in_fence = not in_fence
            out_lines.append(line)
            continue
        if not in_fence:
            match = _LEADING_HASH_RE.match(line)
            if match:
                rest = match.group(2)
                if rest and not rest[0].isspace():
                    line = '\\' + line
        out_lines.append(line)
    return '\n'.join(out_lines)


_PRE_BLOCK_RE = re.compile(r'<pre>\s*(?:<code\b[^>]*>)?(.*?)(?:</code>)?\s*</pre>', re.S)


def harden_code_blocks(html_content):
    """Replaces <pre>/<pre><code> blocks with a <div class="codeblock">,
    turning embedded newlines into explicit <br> tags.

    xhtml2pdf (the PDF renderer) does not implement CSS 'white-space:
    pre-wrap' correctly: it was tested to silently collapse ALL newlines
    (including blank lines between paragraphs) exactly like 'normal' would,
    while 'white-space: pre' preserves newlines but never wraps long lines,
    letting them run off the right edge of the page. Neither single CSS
    value gives both "keep the line breaks" and "wrap long lines".

    The one approach confirmed (by rendering test PDFs) to do both
    correctly in xhtml2pdf is the same one already relied on elsewhere in
    this document via the 'nl2br' extension: real <br> tags plus ordinary
    'white-space: normal' wrapping. This function applies that same trick
    to fenced/indented code blocks, which the 'markdown' library always
    renders as <pre> (optionally with a nested <code>) and leaves as raw
    newlines instead of <br> tags.

    Args:
        html_content (str): HTML produced by markdown.markdown().

    Returns:
        str: The same HTML with every <pre>...</pre> replaced by a
            '<div class="codeblock">' using <br> for line breaks.
    """
    def repl(match):
        inner = match.group(1).rstrip('\n')
        return f'<div class="codeblock">{inner.replace(chr(10), "<br>")}</div>'

    return _PRE_BLOCK_RE.sub(repl, html_content)


def convert_to_html(md_text):
    """Converts Markdown text to HTML string using the 'markdown' library.

    Args:
        md_text (str): The raw Markdown content string to be converted.

    Returns:
        str: The converted HTML string with tables, sane lists, fenced code
            blocks, and new line breaks supported.
    """
    md_text = escape_false_headings(md_text)
    html_content = markdown.markdown(
        md_text, extensions=['tables', 'sane_lists', 'nl2br', 'fenced_code']
    )
    return harden_code_blocks(html_content)


# Printable width of the PDF page (A4 minus the left/right margins set in the
# @page CSS rule below: 210mm page width - 16mm - 16mm margins), converted to
# points. Used to give wide tables absolute (not percentage) column widths.
_MM_TO_PT = 72 / 25.4
PAGE_CONTENT_WIDTH_PT = (210 - 16 - 16) * _MM_TO_PT


def _cell_text_len(cell_html):
    """
    [Utility] Strips all HTML tags from a given cell's content and returns 
    the remaining plaintext length. This is used for calculating content-based column widths.

    Parameters:
        cell_html (str): The raw HTML string of the table cell.

    Returns:
        int: Length of the visible text content after stripping all tags.
    """
    # Regex Explanation: <[^>]+> matches any sequence starting with < and ending with > 
    # (i.e., a complete HTML tag). We substitute these matches with nothing ('') to strip them out.
    return len(re.sub(r'<[^>]+>', '', cell_html).strip())


def _longest_word_len(cell_html: str) -> int:
    """
    [Utility] Calculates the length of the longest continuous, unbreakable word within a table cell's content.

    This calculation is crucial because PDF rendering engines like xhtml2pdf/ReportLab only wrap text 
    at whitespace boundaries. To prevent runtime crashes when calculating column widths (due to 
    negative available width), we must ensure that the minimum calculated column width accommodates 
    the single longest word in any given cell, regardless of surrounding content.

    Parameters:
        cell_html (str): The raw HTML string content of the table cell.

    Returns:
        int: Length of the longest single word found in the cell's text content. Returns 0 if empty.
    """
    # Step 1: Strip all HTML tags to get pure, clean text for analysis using regex substitution.
    text = re.sub(r'<[^>]+>', '', cell_html).strip()
    
    # Step 2: Use regex splitting (re.split) by one or more whitespace characters (\s+)
    # This reliably extracts all word tokens from the continuous text string, handling multiple spaces robustly.
    words = re.split(r'\s+', text)
    
    # Step 3: Calculate and return the maximum length found among all extracted words using a generator expression.
    # Example of advanced Python features: The `max()` function combined with a generator `(len(w) for w in words)`
    # is efficient as it calculates lengths only when needed, avoiding unnecessary list creation.
    return max((len(w) for w in words), default=0)


def shrink_wide_tables(html_content, max_cols_before_shrink=6):
    """Gives explicit per-column widths (and smaller font/padding) to tables
    with many columns, so they fit within the page's printable width.

    When a markdown table has many narrow columns, xhtml2pdf's automatic width
    calculation can produce a negative available width and crash with
    "ValueError: ... negative availWidth ...". This function detects such cases
    and applies explicit per-column widths (in points) plus smaller font/padding
    to prevent the error.

    Widths are set in absolute points (derived from the page's printable width)
    rather than percentages: xhtml2pdf resolves table-cell '%' widths against
    the full page width, not the margin-adjusted printable area, so percentage
    widths cause the last column(s) to spill past the right margin.

    Column widths are proportional to each column's content (the longest cell
    text in that column), with a minimum floor, so text-heavy columns
    (e.g. free-text notes) get more room than short numeric ones.

    Args:
        html_content (str): The HTML content containing one or more tables.
        max_cols_before_shrink (int): Column-count threshold above which a
            table's widths get explicitly resized. Defaults to 6. Tables with
            <= this many columns are returned unchanged.

    Returns:
        str: Modified HTML with adjusted tables that fit within page margins.
    """
    def process_table(match):
        table_html = match.group(0)
        header_row_match = re.search(r'<tr>(.*?)</tr>', table_html, flags=re.S)
        if not header_row_match:
            return table_html
        num_cols = len(re.findall(r'<th\b', header_row_match.group(1)))
        if num_cols <= max_cols_before_shrink:
            return table_html

        extra_cols = num_cols - max_cols_before_shrink
        font_pt = max(6.0, 9.0 - extra_cols * 0.4)
        pad_px = max(2, 6 - extra_cols)
        # Rough Helvetica-Bold average glyph width, used to size each
        # column's hard minimum from its longest unbreakable word.
        char_width_pt = font_pt * 0.62
        pad_pt = pad_px * 0.75 * 2  # px->pt, both sides

        # Longest cell text / longest single word per column, over every row.
        max_len = [0] * num_cols
        max_word_len = [0] * num_cols
        for row_match in re.finditer(r'<tr>(.*?)</tr>', table_html, flags=re.S):
            cells = re.findall(r'<t[dh]\b[^>]*>(.*?)</t[dh]>', row_match.group(1), flags=re.S)
            for col_idx, cell_html in enumerate(cells[:num_cols]):
                max_len[col_idx] = max(max_len[col_idx], _cell_text_len(cell_html))
                max_word_len[col_idx] = max(max_word_len[col_idx], _longest_word_len(cell_html))

        # Slight safety margin (0.97) to absorb border/rounding overhead.
        usable_width_pt = PAGE_CONTENT_WIDTH_PT * 0.97

        # Hard floor: a column can never go below what its longest
        # unbreakable word needs, or xhtml2pdf silently expands it anyway
        # (overflowing the table past the page's right margin).
        hard_min_pt = [
            max(20.0, word_len * char_width_pt + pad_pt)
            for word_len in max_word_len
        ]

        # Weight = content length with a floor, so empty/short columns don't
        # collapse to nothing, and capped so one huge cell can't dominate.
        weights = [max(6, min(length, 40)) for length in max_len]
        total_weight = sum(weights)

        total_hard_min = sum(hard_min_pt)
        if total_hard_min >= usable_width_pt:
            # Content is too wide even at the minimums: fall back to the
            # hard minimums as-is (table will be as tight as it can be).
            col_widths_pt = hard_min_pt
        else:
            extra_width_pt = usable_width_pt - total_hard_min
            col_widths_pt = [
                hard_min_pt[i] + extra_width_pt * weights[i] / total_weight
                for i in range(num_cols)
            ]

        table_html = re.sub(
            r'<table\b[^>]*>',
            f'<table style="width:{PAGE_CONTENT_WIDTH_PT:.1f}pt;font-size:{font_pt:.1f}pt;">',
            table_html,
            count=1,
        )

        # Apply the width to every cell in every row (th AND td), not just
        # the header: xhtml2pdf resets a column's width down to just its
        # padding whenever it meets an EMPTY <td> with no explicit width,
        # silently discarding the width set on the header. Giving every
        # cell its own explicit width avoids that code path entirely.
        def process_row(row_match):
            row_html = row_match.group(0)
            col_idx = {'i': 0}

            def repl_cell(cell_match):
                tag = cell_match.group(1)
                inner = cell_match.group(2)
                i = col_idx['i']
                col_idx['i'] += 1
                width = col_widths_pt[i] if i < len(col_widths_pt) else 20.0
                return (
                    f'<{tag} style="width:{width:.2f}pt;padding:{pad_px}px 4px;">'
                    f'{inner}</{tag}>'
                )

            return re.sub(r'<(th|td)\b[^>]*>(.*?)</\1>', repl_cell, row_html, flags=re.S)

        table_html = re.sub(r'<tr>.*?</tr>', process_row, table_html, flags=re.S)
        return table_html

    return re.sub(r'<table\b[^>]*>.*?</table>', process_table, html_content, flags=re.S)


def html_to_pdf(html_content, output_path):
    """Converts an HTML string to a PDF file using the 'xhtml2pdf' library.

    Args:
        html_content (str): The complete HTML document content, including @page CSS
            with margin settings (210mm - 32mm = 178mm printable width).
        output_path (Path or str): Path where the PDF file will be written.

    Returns:
        bool: True if the PDF was generated successfully, False otherwise.
    """
    try:
        with open(output_path, 'wb') as pdf_file:
            result = pisa.CreatePDF(html_content, dest=pdf_file, encoding='utf-8')
        if result.err:
            print(f"Errore nella generazione del PDF (xhtml2pdf ha segnalato {result.err} problemi)")
            return False
        return True
    except PermissionError as e:
        print(f"Impossibile scrivere su {output_path.name}: {e}")
        return False


def add_page_number_footer(section):
    """Adds a right-aligned page-number field to a section's footer.

    Inserts the real `{ PAGE }` field (same fldChar begin/instrText/separate/end
    pattern as build_toc's TOC field) so Word/LibreOffice compute and keep the
    number in sync as pages are added or removed, from the first page to the
    last.

    Args:
        section: A python-docx Section object (e.g. doc.sections[0]) whose
            footer should get the page number.

    Returns:
        None.
    """
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)

    r_instr = OxmlElement('w:r')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' PAGE '
    r_instr.append(instr_text)

    r_separate = OxmlElement('w:r')
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    r_separate.append(fld_separate)

    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)

    for el in (r_begin, r_instr, r_separate, r_end):
        paragraph._p.append(el)


def build_toc(doc, has_headings, insert_after=None):
    """Builds a native, updatable Table of Contents field for Word documents.

    This inserts the real `{ TOC \\o "1-3" \\h \\z \\u }` field (wrapped in the
    same docPartObj/"Table of Contents" content control Word itself uses),
    with only a placeholder line as its cached result - no hand-built
    hyperlinks or bookmarks. Word/LibreOffice both take ownership of a field
    like this and rebuild its content themselves from the document's
    Heading 1-3 paragraphs whenever the user updates it (F9 / right-click >
    Update Field). Trying to pre-fill the cached result with our own
    hyperlinks was tried and made things worse: both editors detected the
    field, discarded our cached entries as stale, and produced a dead TOC
    that couldn't even be regenerated with F9. Leaving it genuinely empty is
    what makes the F9 update path (confirmed working by the user) kick in.

    Args:
        doc (Document): The python-docx Document object to modify.
        has_headings (bool): Whether the document has any headings (H1-H3).
            If False, no TOC is generated and this function returns immediately.
        insert_after: Optional xml element (e.g. a heading paragraph's ._p) the TOC block
            should be placed right after - typically the document title, so the
            title reads before the TOC instead of being pushed after it. Falls
            back to inserting at the very start of the document when None.

    Returns:
        None. Returns early (without inserting anything) if has_headings is False.
    """
    if not has_headings:
        return

    title_p = doc.add_paragraph()
    title_run = title_p.add_run('Table of Contents')
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    field_p = doc.add_paragraph()
    field_p.paragraph_format.left_indent = Inches(0)

    r_begin = OxmlElement('w:r')
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fld_begin)

    r_instr = OxmlElement('w:r')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr_text)

    r_separate = OxmlElement('w:r')
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    r_separate.append(fld_separate)

    r_placeholder = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    italic = OxmlElement('w:i')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '888888')
    rpr.append(italic)
    rpr.append(color)
    r_placeholder.append(rpr)
    t_placeholder = OxmlElement('w:t')
    t_placeholder.text = (
        "Clic destro qui sopra e scegli \"Aggiorna campo\" (o seleziona il "
        "testo e premi F9) per generare il sommario."
    )
    r_placeholder.append(t_placeholder)

    r_end = OxmlElement('w:r')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fld_end)

    for el in (r_begin, r_instr, r_separate, r_placeholder, r_end):
        field_p._p.append(el)

    doc.add_page_break()

    sdt = OxmlElement('w:sdt')
    sdt_pr = OxmlElement('w:sdtPr')
    doc_part_obj = OxmlElement('w:docPartObj')
    doc_part_gallery = OxmlElement('w:docPartGallery')
    doc_part_gallery.set(qn('w:val'), 'Table of Contents')
    doc_part_unique = OxmlElement('w:docPartUnique')
    doc_part_obj.append(doc_part_gallery)
    doc_part_obj.append(doc_part_unique)
    sdt_pr.append(doc_part_obj)
    sdt.append(sdt_pr)

    sdt_content = OxmlElement('w:sdtContent')
    sdt_content.append(title_p._p)
    sdt_content.append(field_p._p)
    sdt.append(sdt_content)

    body = doc.element.body
    if insert_after is not None and insert_after in list(body):
        body.insert(list(body).index(insert_after) + 1, sdt)
    else:
        body.insert(0, sdt)


_INLINE_MD_PATTERN = re.compile(r'(`[^`]+`|\*\*.+?\*\*|\*.+?\*)')


def strip_inline_markers(text):
    """Removes Markdown inline formatting markers (**bold**, *italic*, `code`)
    from a string, keeping only the wrapped text. Used for H1-H3 headings,
    which are already bold via their Word "Heading" style, so the raw
    markers never show up literally in the heading/TOC.

    Args:
        text (str): Text that may contain inline Markdown formatting markers.

    Returns:
        str: The same text with the markers stripped, wrapped text unchanged.
    """
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text


def add_inline_runs(paragraph, text, code_font='Consolas'):
    """Parses inline Markdown formatting (**bold**, *italic*, `code`) out of
    a line of text and appends it to a Word paragraph as separate runs with
    the matching formatting applied.

    FIX: The previous implementation converted these markers to HTML tags
    (<strong>/<em>/<code>) and then stripped ALL HTML tags with a single
    regex before handing the text to doc.add_paragraph() - so the markers
    were removed but the intended bold/italic/code formatting was silently
    dropped, leaving plain unformatted text (or, for list items/headings/
    table cells, which never went through this pipeline at all, the raw
    **markers** left in place, exactly as reported).

    Args:
        paragraph: The python-docx Paragraph object to append runs to.
        text (str): A line of plain text, potentially containing inline
            Markdown bold/italic/code markers.
        code_font (str): Font name applied to inline code runs. Defaults to
            'Consolas' to match the fenced code block styling.

    Returns:
        None.
    """
    for part in _INLINE_MD_PATTERN.split(text):
        if not part:
            continue
        if part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = code_font
        elif part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            paragraph.add_run(part)


def create_word_from_md(md_text, output_path, doc_title=None, doc_subtitle=None):
    """Creates a Microsoft Word document (.docx) from Markdown source with automatic
    Table of Contents (TOC) generation. Supports all standard Markdown syntax
    including headings, lists, tables, images, links, and code blocks.

    TOC Features:
    - Automatically numbered headings (1, 1.1, 1.1.1 format): H1 is the
      chapter counter (1, 2, 3...), H2 nests under its enclosing H1 (1.1,
      1.2...), and H3 nests under its enclosing H2 (1.1.1, 1.1.2...). The
      markdown content itself has no separate, un-numbered "document title"
      heading - that comes only from doc_title/doc_subtitle (front matter),
      if given.
    - Clickable TOC with hyperlinks to each section
    - Hierarchical indentation matching heading levels

    Supported Markdown features:
    - Headings: # h1, ## h2, ### h3, #### h4, ##### h5, ###### h6
    - Unordered lists (- or *)
    - Ordered lists (1., 2., ...)
    - Bold (**text**) and Italic (*text*)
    - Inline code (`code`)
    - Tables with header rows
    - Links [text](url)
    - Images ![alt](path)

    Document Format:
    - Page: ISO A4 (21.0 x 29.7 cm), portrait - overrides python-docx's
      built-in default of US Letter (21.59 x 27.94 cm).
    - Margins: left 2.5 cm / right 2.10 cm are set explicitly (an
      asymmetric pair matching the printable width used by the HTML/PDF
      outputs). Top and bottom margins (2.54 cm / 1 inch each) and the
      header/footer distance (1.27 cm / 0.5 inch) are left at python-docx's
      built-in template defaults - only left/right are overridden here.
    - Fonts: no custom font or theme is defined in code, so text renders in
      whatever fonts the built-in python-docx template's Office theme
      assigns to each style. Body text (the Normal style) uses the theme's
      minor font, Cambria; the Title style and Heading 1-3 use the theme's
      major font, Calibri - a humanist sans-serif (its letterforms follow
      classic handwritten proportions rather than the rigid geometric grid
      of a grotesque sans like Arial/Helvetica). The one explicit override
      is fenced code blocks, forced to the monospace font Consolas.
    - Footer: every page gets a right-aligned, auto-updating page-number
      field, added via add_page_number_footer() (see its docstring for how
      the field mechanism works).

    Args:
        md_text (str): The full Markdown content as a string
        output_path: Path where the .docx file will be saved
        doc_title (str, optional): Optional document title (from front matter),
            rendered with the big built-in "Title" style before the TOC. Defaults to None.
        doc_subtitle (str, optional): Optional subtitle (from front matter), rendered
            right under doc_title. Ignored if doc_title is not given. Defaults to None.

    Returns:
        bool: True if document was saved successfully, False on error.
    """
    def split_table_row(row_line):
        # Helper: splits a markdown table row by '|' into individual cell content
        parts = row_line.split('|')
        
        # Remove empty first column (common pattern in markdown tables)
        if parts and parts[0].strip() == '':
            parts = parts[1:]
        
        # Remove empty last column if present
        if parts and parts[-1].strip() == '':
            parts = parts[:-1]
        
        return [p.strip() for p in parts]

    try:
        doc = Document()

        section = doc.sections[0]
        # A4 portrait (overrides python-docx's built-in US Letter default);
        # only left/right margins are set here - top/bottom margin and the
        # header/footer distance stay at the built-in template defaults
        # (2.54 cm / 1 inch top+bottom, 1.27 cm / 0.5 inch header/footer).
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.10)
        add_page_number_footer(section)

        h_counters = [0, 0, 0]
        has_headings = False
        title_element = None  # xml element the TOC is inserted right after

        if doc_title:
            heading = doc.add_heading(doc_title, level=0)
            heading.paragraph_format.space_after = Pt(4 if doc_subtitle else 0)
            title_element = heading._p
            if doc_subtitle:
                subtitle_p = doc.add_paragraph()
                subtitle_run = subtitle_p.add_run(doc_subtitle)
                subtitle_run.font.size = Pt(14)
                subtitle_run.font.italic = True
                subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                subtitle_p.paragraph_format.space_after = Pt(0)
                title_element = subtitle_p._p

        lines = md_text.split('\n')
        num_lines = len(lines)
        i = 0

        while i < num_lines:
            line_stripped = lines[i]
            line = line_stripped.strip()

            if not line:
                # Skip empty lines (no paragraph needed)
                i += 1
                continue

            # ── Fenced Code Blocks (```...``` or ~~~...~~~) ──
            # Render the block's lines verbatim as monospace paragraphs and
            # skip normal markdown parsing inside it - otherwise a Python/SQL
            # comment like "# nota" sitting inside the fence would be
            # mistaken for a real heading.
            fence_match = re.match(r'^(```|~~~)', line)
            if fence_match:
                fence = fence_match.group(1)
                j = i + 1
                while j < num_lines and not lines[j].strip().startswith(fence):
                    p = doc.add_paragraph(lines[j])
                    for run in p.runs:
                        run.font.name = 'Consolas'
                    j += 1
                i = j + 1  # skip past the closing fence line
                continue

            # ── Markdown Tables Detection ──
            # Detect table syntax: row containing '|' with a separator header row above it
            if '|' in line and not line.startswith('#'):
                # Look ahead at the next line to find the separator (header divider) pattern
                next_line = lines[i + 1].strip() if i + 1 < num_lines else ''
                sep_cells = split_table_row(next_line) if next_line else []
                
                # Check if this is a markdown table separator row (e.g., "|-|---|---:")
                is_separator = bool(sep_cells) and all(
                    re.match(r'^:?-+:?$', c) for c in sep_cells if c != ''
                ) and any(c != '' for c in sep_cells)

                if is_separator:
                    # This is a header separator, now extract the header row above it
                    header_cells = split_table_row(line)
                    num_cols = len(header_cells)

                    # Collect all data rows below until we hit another table section or non-table line
                    body_rows = []
                    j = i + 2
                    while j < num_lines and '|' in lines[j].strip():
                        body_rows.append(split_table_row(lines[j].strip()))
                        j += 1

                    # Create the Word table with header and body rows
                    table = doc.add_table(rows=1 + len(body_rows), cols=num_cols)
                    table.style = 'Table Grid'

                    # Fill in header cells (first row)
                    for k, cell_text in enumerate(header_cells[:num_cols]):
                        add_inline_runs(table.rows[0].cells[k].paragraphs[0], cell_text)

                    # Fill in body cells (data rows)
                    for r, cells in enumerate(body_rows, start=1):
                        for k, cell_text in enumerate(cells[:num_cols]):
                            add_inline_runs(table.rows[r].cells[k].paragraphs[0], cell_text)

                    i = j  # Move to next section
                    continue

            # ── Images ![alt](path) ──
            # Detect markdown image syntax and convert to plain text description
            img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line_stripped)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                p = doc.add_paragraph(f"Image: {alt_text} (path: {img_path})")
                i += 1
                continue

            # ── Links [text](url) ──
            # Detect markdown link syntax and convert to clickable-style text
            link_match = re.match(r'^\[([^\]]*)\]\(([^)]+)\)$', line_stripped)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)
                p = doc.add_paragraph(f"[{link_text}]({link_url})")
                i += 1
                continue
            
            # ── Headings (h1-h6) ──
            # Detect markdown headings (# to ######) and convert to Word headings with bookmarks
            title_match = re.match(r'^(#{1,6})\s+(.+)$', line_stripped)
            if title_match:
                level = len(title_match.group(1))
                title_text = title_match.group(2).strip()
                
                # Remove any [text](url) links from the heading, keeping only the text
                title_clean = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', title_text)
                
                if level in (1, 2, 3):
                    # H1 is the chapter counter (1, 2, 3...); H2 nests under
                    # its enclosing H1 (1.1, 1.2...); H3 nests under its
                    # enclosing H2 (1.1.1, 1.1.2...). If a document jumps
                    # straight to H2/H3 without a preceding H1, the missing
                    # ancestor levels are treated as chapter/section 1 so the
                    # numbering never shows a leading "0".
                    if level == 1:
                        h_counters[0] += 1
                        h_counters[1] = 0
                        h_counters[2] = 0
                        numbering = f"{h_counters[0]}"
                    elif level == 2:
                        if h_counters[0] == 0:
                            h_counters[0] = 1
                        h_counters[1] += 1
                        h_counters[2] = 0
                        numbering = f"{h_counters[0]}.{h_counters[1]}"
                    else:
                        if h_counters[0] == 0:
                            h_counters[0] = 1
                        if h_counters[1] == 0:
                            h_counters[1] = 1
                        h_counters[2] += 1
                        numbering = f"{h_counters[0]}.{h_counters[1]}.{h_counters[2]}"

                    numbered_text = f"{numbering} {strip_inline_markers(title_clean)}"
                    heading = doc.add_heading(numbered_text, level=level)
                    has_headings = True
                    if title_element is None:
                        title_element = heading._p
                elif level == 4:
                    # Level 4 headings are too deep for TOC, render as custom paragraph
                    p = doc.add_paragraph()
                    add_inline_runs(p, title_clean)
                    # Match the theme's major (heading) font used by the real
                    # H1-H3 headings and the Title style, since H4 is styled
                    # to look like a heading even though it's not TOC-tracked.
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(14)
                elif level == 5:
                    # Level 5 headings are very small detail headers
                    p = doc.add_paragraph()
                    add_inline_runs(p, title_clean)
                    for run in p.runs:
                        run.font.size = Pt(12)
                else:  # h6
                    # Level 6 headings are fine print / footer-level text
                    p = doc.add_paragraph()
                    add_inline_runs(p, title_clean)
                    for run in p.runs:
                        run.font.size = Pt(10)

                i += 1
                continue

            # ── Unordered Lists (-, *, +) ──
            # Detect bullet point lists and add them as paragraphs
            ul_match = re.match(r'^(\s*[-*+]\s+)(.+)$', line_stripped)
            if ul_match:
                list_item_text = ul_match.group(2).strip()
                p = doc.add_paragraph()
                add_inline_runs(p, list_item_text)
                i += 1
                continue

            # ── Ordered Lists (1., 2., 3...) ──
            # Detect numbered lists and preserve the numbering
            ol_match = re.match(r'^(\d+)\.\s+(.+)$', line_stripped)
            if ol_match:
                list_num = ol_match.group(1)
                list_item_text = ol_match.group(2).strip()
                p = doc.add_paragraph()
                p.add_run(f"{list_num}. ")
                add_inline_runs(p, list_item_text)
                i += 1
                continue

            # ── Blockquotes (") ──
            # Detect quoted text and add with Quote style in Word
            quote_match = re.match(r'^"(.+)"$', line_stripped)
            if quote_match:
                quote_text = quote_match.group(1).replace('"', '"')
                p = doc.add_paragraph(style='Quote')
                p.add_run('" ')
                add_inline_runs(p, quote_text)
                p.add_run('"')
                i += 1
                continue

            # ── Regular Paragraph Text ──
            # Add remaining text as a regular paragraph (not heading/list/quote/etc.),
            # with inline **bold**/*italic*/`code` rendered as real Word formatting.
            if line:
                p = doc.add_paragraph()
                add_inline_runs(p, line)

            i += 1

        build_toc(doc, has_headings, insert_after=title_element)

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Errore nella creazione del file DOCX: {e}")
        return False


# ── CSS stylesheet for HTML/PDF rendering ──
# Includes TOC styling and anchor links for clickable navigation
CSS = """
@page { size: A4; margin: 18mm 16mm 16mm 16mm; }
* { box-sizing: border-box; }

body { 
    font-family: 'Lato','Segoe UI',Helvetica,Arial,sans-serif; 
    font-size: 10pt; 
    line-height: 1.55; 
    color: #2b2b2b; 
    margin: 0; 
}

h1, h2, h3, h4, h5, h6 {
    margin-top: 0;
    page-break-after: avoid;
}

/* Document title/subtitle (from front matter, shown above the TOC) */
.doc-title {
    font-size: 26pt;
    font-weight: 700;
    color: #1e3c72;
    margin: 0 0 4px 0;
    letter-spacing: .3px;
}
.doc-subtitle {
    font-size: 13pt;
    color: #555;
    font-style: italic;
    margin: 0 0 24px 0;
}

/* Table of Contents container */
#toc { 
    background: #f8f9fa; 
    padding: 15px 20px; 
    border-radius: 6px; 
    margin-top: 0; 
    margin-bottom: 15px;
    border-left: 4px solid #1e3c72;
}

#toc h4 { 
    margin: 0 0 10px 0; 
    color: #1e3c72; 
    font-size: 12pt;
    font-weight: bold;
}

/* TOC navigation links */
#toc nav ul { 
    list-style: none; 
    padding: 0; 
    margin: 0;
}

#toc nav li {
    margin: 4px 0;
    padding-left: 15px;
    position: relative;
}

#toc nav li::before {
    content: "-";
    position: absolute;
    left: 0;
    font-size: 9.5pt;
}

/* Hierarchical indentation, mirrors the Word TOC (0 / 0.25in / 0.5in) */
#toc nav li.toc-l1 { margin-left: 0; font-weight: 600; }
#toc nav li.toc-l2 { margin-left: 18px; }
#toc nav li.toc-l3 { margin-left: 36px; font-size: 9pt; }

/* TOC links - clickable and styled */
#toc nav a {
    color: #1e3c72;
    text-decoration: none;
    font-weight: 500;
    font-size: 9.5pt;
    padding: 2px 4px;
    border-radius: 3px;
    transition: all 0.2s ease;
}

#toc nav a:hover { 
    background: #e8f0fe; 
    text-decoration: underline;
}

/* Headings with anchor links */
h1, h2, h3, h4, h5, h6 { 
    scroll-margin-top: 150px;
}

h1 { font-size: 20pt; color: #1e3c72; margin: 0 0 8px 0; letter-spacing: .3px; }
h2 { font-size: 12.5pt; color: #1e3c72; margin: 22px 0 8px 0; padding-bottom: 4px; border-bottom: 2px solid #1e3c72; }
h3 { font-size: 10.8pt; color: #2a5298; margin: 15px 0 5px 0; }
h4 { font-size: 10pt; color: #444; margin: 12px 0 4px 0; }
h5 { font-size: 9.6pt; color: #555; margin: 10px 0 3px 0; font-style: italic; }

p { margin: 0 0 7px 0; }

ul, ol { margin: 4px 0 8px 0; padding-left: 20px; }
li { margin-bottom: 3px; }

strong { color: #1e3c72; font-weight: bold; }
em { color: #555; font-style: italic; }

code {
    background: #eef1f5;
    color: #1e3c72;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Consolas','Courier New',monospace;
    font-size: 9pt;
}

/* Fenced/indented code blocks: harden_code_blocks() rewrites the
   markdown library's <pre> output into this div (see its docstring for
   why - xhtml2pdf's 'white-space: pre-wrap' silently drops line breaks). */
.codeblock {
    background: #2b2f3a;
    color: #e8e8e8;
    border-left: 4px solid #1e3c72;
    padding: 10px 12px;
    margin: 8px 0 12px 0;
    font-family: 'Consolas','Courier New',monospace;
    font-size: 9pt;
    white-space: normal;
    word-wrap: break-word;
    overflow-wrap: break-word;
}

table {
    width: 100%;
    table-layout: fixed;
    border-collapse: collapse;
    margin: 8px 0 12px 0;
    font-size: 9pt;
}
th {
    background: #eef3fb;
    color: #1e3c72;
    text-align: left;
    padding: 6px 8px;
    border: 1px solid #cdd9ec;
    font-weight: bold;
    word-wrap: break-word;
    overflow-wrap: break-word;
}
td {
    padding: 6px 8px;
    border: 1px solid #dde5f0;
    vertical-align: top;
    word-wrap: break-word;
    overflow-wrap: break-word;
}

blockquote { 
    margin: 7px 0 10px 0; 
    padding: 8px 14px; 
    background: #f5f7fa; 
    border-left: 4px solid #1e3c72; 
    font-size: 9.6pt; 
    page-break-inside: avoid;
}
blockquote p { margin: 0; }

hr { border: none; border-top: 1px solid #e2e2e2; margin: 15px 0; }

a { color: #1e3c72; text-decoration: underline; }
img { max-width: 100%; height: auto; }
"""

def slugify(text):
    """Converts a string to a URL-safe anchor ID for headings.

    Performs NFKD normalization, strips non-word characters, collapses whitespace,
    and lowercases the result. Returns an empty string or 'section' if input is empty.

    FIX: Previously, TOC links calculated the anchor from text using a
    different rule than the one used to generate IDs on headings, so they
    never matched. Now both <a href="#..."> from TOC and heading IDs pass
    through this single function (with duplicate handling done by caller),
    so they always stay aligned.

    Args:
        text (str): The heading text to convert into a slug.

    Returns:
        str: A URL-safe anchor ID (lowercase, hyphen-separated words). Empty
            string or 'section' if the input was empty after normalization.
    """
    text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('ascii')
    text = re.sub(r'[^\w\s-]', '', text).strip().lower()
    text = re.sub(r'[-\s]+', '-', text)
    return text or 'section'


def main():
    # ── Interactive menu for choosing the output format ──
    print("\n" + "=" * 40)
    print("=== Conversione file Markdown (con TOC per HTML/PDF) ===")
    print("=" * 40)
    print("H = Solo HTML (.html)")
    print("P = Solo PDF (.pdf)  <-- DEFAULT")
    print("W = Solo Word (.docx)")
    print("A = Tutti e tre i formati (HTML, PDF, DOCX)")
    print("=" * 40)

    choice = input("\nScegli il formato: H | P | W | A? ").strip().upper()

    # Fallback: if the user doesn't select anything, PDF is used as the default
    if not choice or choice == "":
        print("Nessuna selezione - uso il default: PDF (.pdf)")
        choice = "P"

    if choice not in ['H', 'P', 'W', 'A']:
        print("Selezione non valida. Uscita dal programma.")
        sys.exit(1)

    # ── Input file ──
    arg = sys.argv[1] if len(sys.argv) > 1 else 'text_to_convert.md'
    src = Path(arg)
    if not src.is_absolute():
        src = SCRIPT_DIR / src

    if not src.exists():
        print(f"File non trovato: {src}")
        sys.exit(1)

    # ── Output file name ──
    default_stem = src.with_suffix('').name
    output_name = input(f"\nNome base del file di output [{default_stem}]: ").strip()
    stem = Path(output_name) if output_name else src.with_suffix('')

    # ── Determine output files and request confirmation ──
    # These paths must match exactly the ones used further down
    # to save the files (stem.with_suffix(...)), otherwise the existence
    # check looks in the wrong folder and fails to detect files to overwrite.
    target_files = []

    if choice in ['H', 'A']:
        target_files.append(stem.with_suffix('.html'))
    if choice in ['P', 'A']:
        target_files.append(stem.with_suffix('.pdf'))
    if choice in ['W', 'A']:
        target_files.append(stem.with_suffix('.docx'))

    if not target_files:
        # Case where no output is selected (shouldn't happen given the earlier checks, but it's a safety measure)
        print("\nNessun formato di output selezionato. Uscita dal programma.")
        sys.exit(0)

    # Overwrite confirmation should only be asked for files that already
    # exist with that name - if none of the target files exist, proceed
    # directly without interrupting the flow with an unnecessary question.
    existing_files = [f for f in target_files if f.exists()]

    if existing_files:
        print("\n" + "=" * 40)
        print("=== ATTENZIONE: FILE DI OUTPUT GIA' ESISTENTI === ")
        print("I file seguenti esistono gia' e verranno sovrascritti:")
        for i, f in enumerate(existing_files):
            print(f"{i+1}. {f.name}")
        print("=" * 40)

        confirmation = input("\nVuoi procedere con la sovrascrittura dei file mostrati?(S/N) [Default N]: ").strip().upper()

        if confirmation == 'S':
            # Consent received, continue execution
            pass
        else:
            print("\nOk allora non devo fare nulla. Uscita dal programma.")
            sys.exit(0) # Terminate the script without saving files


    # ── Read Markdown content ──
    print("\nLettura del file Markdown in corso...")
    try:
        md_text = open(src, encoding='utf-8-sig').read()
    except UnicodeDecodeError:
        print("Attenzione: il file non e' UTF-8, provo con la codifica Windows-1252 (cp1252)...")
        md_text = open(src, encoding='cp1252').read()
    # The document title/subtitle shown above the TOC comes from the front
    # matter (if any), not from the first '#' heading, which is now numbered
    # and indexed like any other chapter - so the metadata must be read before
    # strip_frontmatter discards the block.
    frontmatter_fields = extract_frontmatter_fields(md_text)
    doc_title = frontmatter_fields.get('title')
    doc_subtitle = frontmatter_fields.get('subtitle')
    md_text = strip_frontmatter(md_text)
    # Remove any hand-written TOC in the source markdown: HTML, PDF and DOCX
    # all generate their own automatic TOC further down, so the original
    # one would just be duplicated.
    md_text = strip_md_toc(md_text)
    print(f"Letti: {len(md_text)} caratteri.")

    # ── Convert Markdown to HTML (required for HTML/PDF output) ──
    needs_html = choice in ['H', 'P', 'A']
    body_html = None
    if needs_html and MARKDOWN_AVAILABLE:
        body_html = convert_to_html(md_text)
        body_html = shrink_wide_tables(body_html)


    # ── Assemble the HTML document with a hierarchical TOC (mirrors the Word TOC) ──
    html_doc = None
    if body_html is not None:
        # Numbers h1-h3 (1, 1.1, 1.1.1 ...) like the Word version, gives every heading
        # a unique anchor id, and collects h1-h3 entries for the TOC.
        h_counters = [0, 0, 0]
        used_ids = {}
        toc_entries = []  # (level, numbered_text, anchor_id)

        def add_heading_id(match_obj):
            level = int(match_obj.group(1))
            inner_html = match_obj.group(2)

            text_plain = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', inner_html)
            text_plain = re.sub(r'<[^>]+>', '', text_plain)
            text_plain = ' '.join(text_plain.split())

            display_html = inner_html
            if level in (1, 2, 3):
                # H1 is the chapter counter (1, 2, 3...); H2 nests under its
                # enclosing H1 (1.1, 1.2...); H3 nests under its enclosing H2
                # (1.1.1, 1.1.2...). Missing ancestor levels (e.g. a document
                # jumping straight to H2/H3) are treated as chapter/section 1
                # so numbering never shows a leading "0".
                if level == 1:
                    h_counters[0] += 1
                    h_counters[1] = 0
                    h_counters[2] = 0
                    numbering = f"{h_counters[0]}"
                elif level == 2:
                    if h_counters[0] == 0:
                        h_counters[0] = 1
                    h_counters[1] += 1
                    h_counters[2] = 0
                    numbering = f"{h_counters[0]}.{h_counters[1]}"
                else:
                    if h_counters[0] == 0:
                        h_counters[0] = 1
                    if h_counters[1] == 0:
                        h_counters[1] = 1
                    h_counters[2] += 1
                    numbering = f"{h_counters[0]}.{h_counters[1]}.{h_counters[2]}"
                display_html = f"{numbering} {inner_html}"
                text_plain = f"{numbering} {text_plain}"

            base_id = slugify(text_plain)
            seen = used_ids.get(base_id, 0)
            used_ids[base_id] = seen + 1
            anchor_id = base_id if seen == 0 else f"{base_id}-{seen}"

            if level in (1, 2, 3):
                toc_entries.append((level, display_html, anchor_id))

            # FIX: PDF links were broken because xhtml2pdf (reportlab engine) ignores
            # the name="..." attribute set directly on h1..h6 - it only recognizes a
            # standalone <a name="..."> tag as an internal link target, and discards
            # empty ones during rendering (known bug). So we prepend the anchor tag
            # with a non-breaking space inside: id="..." is used as the target for
            # links in HTML export, while <a name="..."> makes PDF links clickable.
            return (
                f'<a name="{anchor_id}">&nbsp;</a>'
                f'<h{level} id="{anchor_id}" class="anchor">{display_html}</h{level}>'
            )

        html_with_ids = re.sub(
            r'<h([1-6])\b[^>]*>(.*?)</h\1>',
            add_heading_id,
            body_html,
            flags=re.S,
        )

        level_class = {1: 'toc-l1', 2: 'toc-l2', 3: 'toc-l3'}
        toc_links_html = '\n'.join(
            f'<li class="{level_class[level]}"><a href="#{anchor_id}">{text}</a></li>'
            for level, text, anchor_id in toc_entries
        )

        title_box_html = ''
        if doc_title:
            subtitle_html = f'<div class="doc-subtitle">{doc_subtitle}</div>' if doc_subtitle else ''
            title_box_html = f'<h1 class="doc-title">{doc_title}</h1>{subtitle_html}'

        html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>Converted Document</title>
<style>{CSS}</style></head>
<body>

{title_box_html}

<div id="toc">
<h4>Table of Contents</h4>
<nav>
<ul>
{toc_links_html}
</ul>
</nav>
</div>

{html_with_ids}</body>
</html>"""

    # ── Generate the files based on the user's choice ──
    if choice in ['H', 'A']:
        if html_doc is not None:
            html_path = stem.with_suffix('.html')
            html_path.write_text(html_doc, encoding='utf-8')
            print(f"Generato: {html_path}")
        else:
            _missing('markdown')
            print("Generazione HTML saltata.")

    if choice in ['P', 'A']:
        if html_doc is None:
            _missing('markdown')
            print("Generazione PDF saltata.")
        elif not XHTML2PDF_AVAILABLE:
            _missing('xhtml2pdf')
            print("Generazione PDF saltata.")
        else:
            pdf_path = stem.with_suffix('.pdf')
            if html_to_pdf(html_doc, pdf_path):
                print(f"Generato: {pdf_path}")

    # The Word conversion uses the already-processed text (TOC removed from the source markdown)
    if choice in ['W', 'A']:
        if DOCX_AVAILABLE:
            word_path = stem.with_suffix('.docx')
            if create_word_from_md(md_text, word_path, doc_title=doc_title, doc_subtitle=doc_subtitle):
                print(f"Generato: {word_path}")
        else:
            _missing('python-docx')
            print("Generazione DOCX saltata.")

    print("\n" + "=" * 40)
    print("Conversione completata!")
    print("=" * 40)


if __name__ == "__main__":
    main()
